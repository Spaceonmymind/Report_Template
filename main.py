import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QVBoxLayout, QWidget, QFileDialog
import pandas as pd
from docx import Document

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Генератор Отчетов о Нарушениях")
        layout = QVBoxLayout()

        btn_load = QPushButton("Загрузить данные")
        btn_load.clicked.connect(self.load_data)
        layout.addWidget(btn_load)

        btn_report = QPushButton("Создать отчет")
        btn_report.clicked.connect(self.create_report)
        layout.addWidget(btn_report)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.data = None

    def load_data(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "Выберите файл данных", "", "Excel Files (*.xls *.xlsx);;CSV Files (*.csv)")
        if file_name:
            if file_name.endswith('.csv'):
                self.data = pd.read_csv(file_name)
            elif file_name.endswith('.xlsx'):
                self.data = pd.read_excel(file_name, engine='openpyxl')
            elif file_name.endswith('.xls'):
                self.data = pd.read_excel(file_name, engine='xlrd')
            print("Данные загружены")

    def load_data(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "Выберите файл данных", "",
                                                   "Excel Files (*.xls *.xlsx);;CSV Files (*.csv)")
        if file_name:
            if file_name.endswith('.csv'):
                self.data = pd.read_csv(file_name)
            elif file_name.endswith('.xlsx'):
                self.data = pd.read_excel(file_name, engine='openpyxl')
            elif file_name.endswith('.xls'):
                self.data = pd.read_excel(file_name, engine='xlrd')
            print("Данные загружены")
            print("Столбцы в данных:", self.data.columns.tolist())  # Выводим список столбцов для проверки

    def create_report(self):
        if self.data is not None:
            document = Document()
            document.add_heading('Утверждаю:', level=2)

            # Добавляем вводный блок
            document.add_paragraph(
                "Заместитель директора Департамента\n"
                "экономического развития Курганской области\n\n"
                "____________________  И.В. Смирных \n"
                "«____» ___________________ 2023 г.\n\n\n"
                "ОТЧЕТ\n"
                "о результатах мероприятия по контролю без взаимодействия с юридическими лицами, индивидуальными предпринимателями\n\n"
                "«05» мая 2023 года\n\n"
            )

            # Динамическое добавление данных из таблицы
            conducted_date = self.data['Дата розничной продажи (чека)'].iloc[0]
            document.add_paragraph(
                f"На основании задания начальника отдела развития потребительского рынка - управления развития "
                f"рыночной инфраструктуры Департамента экономического развития Курганской области Машуковой О.А. "
                f"от 04 мая 2023 года № 9 проведен анализ о своевременном представлении и достоверности сведений, "
                f"указанных в декларации об объеме розничной продажи алкогольной и спиртосодержащей продукции за "
                f"I квартал 2023 г. Мероприятия были проведены {conducted_date}.\n\n"
                f"Бузмаков Александр Александрович – главный специалист сектора по лицензированию управления развития "
                f"рыночной инфраструктуры Департамента экономического развития Курганской области.\n\n"
            )

            # Пример добавления данных по нарушениям
            for index, row in self.data.iterrows():
                product_name = row['Наименование продукции']
                product_volume = row['Кол-во (шт)']
                product_price = row['Цена\n(руб)']
                document.add_paragraph(
                    f"Было установлено, что продукция {product_name} была продана в количестве {product_volume} по цене {product_price} рублей каждая.\n"
                )

            # Добавление заключительного блока
            document.add_paragraph(
                "Выводы, предложения: ООО «РОМАШКА» осуществляло розничную продажу алкогольной продукции в I квартале 2023 г., "
                "сдало декларацию по форме 37 за I квартал 2023 г. позднее 20-го числа месяца, следующего за отчетным кварталом, "
                "чем нарушило требования пункта 13 Приказа Росалкогольрегулирования.\n\n"
                "Мотивированное представление по результатам анализа мероприятий по контролю без взаимодействия: вызвать представителя "
                "ООО «РОМАШКА» для составления протокола об административном правонарушении по ст.15.13 КоАП РФ, извещение о явке "
                "направить на адрес электронной почты: su4588@mail.ru.\n\n"
            )

            # Сохранение отчета
            document.save('Отчет_о_нарушениях.docx')
            print("Отчет создан")
        else:
            print("Данные не загружены")


app = QApplication(sys.argv)
window = MainWindow()
window.show()
sys.exit(app.exec_())